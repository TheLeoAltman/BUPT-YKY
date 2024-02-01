#coding=utf-8
import os.path

import docx
from docx import Document

class doc:
    def __init__(self, path, filename):
        def doc_to_txt(doc_file, txt_file):
            doc = docx.Document(doc_file)
            paragraphs = [p.text for p in doc.paragraphs]
            with open(txt_file, "w", encoding='utf-8') as f:
                f.write("\n".join(paragraphs))

        doc_file = path + '/' + filename
        length = len(filename)
        for i in range(length - 1, -1, -1):
            if filename[i] == '.':
                break
        name = 'trans/'
        for j in range(0, i + 1):
            name = name + filename[j]
        name += 'txt'
        doc_to_txt(doc_file, name)
class form:
    def __init__(self, path, filename):
        path = os.path.abspath(path)
        in_path = path + '/' +filename
        length = len(filename)
        for i in range(length - 1, -1, -1):
            if filename[i] == '.':
                break
        out_path = os.path.abspath('trans')
        out_path += '/'
        for j in range(0, i):
            out_path = out_path + filename[j]
        out_path += '(form).txt'
        form = Document(in_path)

        tables = form.tables
        print(in_path)
        print(out_path)
        for i in range(len(tables)):
            data = ''
            tb = tables[i]
            # 获取表格的行
            tb_rows = tb.rows
            # 读取每一行内容
            for i in range(len(tb_rows)):
                row_data = []
                row_cells = tb_rows[i].cells
                # 读取每一行单元格内容
                for cell in row_cells:
                    # 单元格内容
                    row_data.append(cell.text)
                    data = '|'.join(row_data)
                data += '\n'
                with open(out_path, 'a', encoding='utf-8') as f:
                    f.write(data)